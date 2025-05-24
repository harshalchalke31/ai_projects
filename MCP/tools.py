from langchain.tools import Tool
from langchain_community.tools import WikipediaQueryRun, DuckDuckGoSearchRun, ArxivQueryRun
from langchain_community.utilities import WikipediaAPIWrapper
from datetime import datetime
from docx import Document
import json

search = DuckDuckGoSearchRun()
search_tool = Tool(
    name="DuckDuckGoSearch",
    func=search.run,
    description="A tool to search the web using DuckDuckGo. Use this tool when you need to find information on the web.",
)

api_wrapper = WikipediaAPIWrapper(top_k_results=3,doc_content_chars_max=200)
wiki_tool = Tool(
    name="WikipediaSearch",
    func=WikipediaQueryRun(api_wrapper=api_wrapper).run,
    description="A tool to search Wikipedia. Use this tool when you need to find information on Wikipedia.",
)

from datetime import datetime
from docx import Document
import json

def save_to_docx(data: str, filename: str = 'research_output'):
    """
    Save structured JSON data to a .docx file in a readable format.
    """
    try:
        parsed = json.loads(data)
    except json.JSONDecodeError:
        # Fallback if input is plain text
        parsed = {"Raw Output": data}

    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    full_filename = f"{filename}_{timestamp}.docx"

    doc = Document()
    doc.add_heading('Research Output', level=1)
    doc.add_paragraph(f"Timestamp: {timestamp}\n")

    # Format each field nicely
    if "topic" in parsed:
        doc.add_heading("Topic", level=2)
        doc.add_paragraph(parsed["topic"])

    if "answer" in parsed:
        doc.add_heading("Answer", level=2)
        doc.add_paragraph(parsed["answer"])

    if "summary" in parsed:
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(parsed["summary"])

    if "sources" in parsed and isinstance(parsed["sources"], list):
        doc.add_heading("Sources", level=2)
        for source in parsed["sources"]:
            doc.add_paragraph(source, style='ListBullet')

    if "tools" in parsed and isinstance(parsed["tools"], list):
        doc.add_heading("Tools Used", level=2)
        for tool in parsed["tools"]:
            doc.add_paragraph(tool, style='ListBullet')

    # Fallback in case it's just raw data
    if "Raw Output" in parsed:
        doc.add_paragraph(parsed["Raw Output"])

    doc.save(full_filename)
    return f"Data saved to {full_filename}"



save_tool = Tool(
    name="SaveToDocx",
    func=save_to_docx,
    description="A tool to save the data to a docx file. Use this tool when you need to save the data to a file.",
)

arxiv = ArxivQueryRun()
arxiv_tool = Tool(
    name="ArxivSearch",
    func=arxiv.run,
    description="A tool to search Arxiv. Use this tool when you need to find research papers on Arxiv.",
)